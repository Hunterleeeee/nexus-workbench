"""文档工厂领域。

拆自 app.py（2026-08-14 第十九批）。包含: 文档生成/校验/评审/再生成/交付。
仍在 app.py 的领域函数经 _app_call 运行时转发。
"""
from __future__ import annotations

import asyncio
import os
import shutil
import subprocess
import uuid
import httpx
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from fastapi import File, HTTPException, UploadFile
from pydantic import BaseModel, Field

from .agent_runs import add_agent_run_event, create_agent_run_record, get_agent_run, update_agent_run_record
from .core import OUTPUTS_DIR, WORKBENCH_VERSION, clip, clip_for_llm, log, now_iso
from .db import db_connection
from .instance import app
from .knowledge import extract_upload_text, knowledge_tokens, write_knowledge_note
from .llm import call_llm, llm_settings
from .notifications import create_notification_record
from .projects import agent_display_name, load_projects


def _DOC_FACTORY_TEMPLATES() -> dict[str, Any]:
    """运行时读 app.DOC_FACTORY_TEMPLATES——测试 patch app.DOC_FACTORY_TEMPLATES 时生效。"""
    import app as _app

    return _app.DOC_FACTORY_TEMPLATES


def _app_call(fn_name: str, *args: Any, **kwargs: Any) -> Any:
    """通过 app 命名空间调用仍在 app.py 的领域函数——测试 patch app.X 时能生效。"""
    import app as _app

    return getattr(_app, fn_name)(*args, **kwargs)


class DocumentFactoryRequest(BaseModel):
    title: str = Field(default="未命名产物", max_length=160)
    source_text: str = Field(default="", max_length=100_000)
    instruction: str = Field(default="", max_length=4_000)
    template: str = Field(default="general_report", max_length=50)
    source_name: str = Field(default="粘贴材料", max_length=240)
    artifact_ids: list[int] = Field(default_factory=list, max_length=12)
    revision_focus: list[str] = Field(default_factory=list, max_length=8)
    acceptance_criteria: list[str] = Field(default_factory=list, max_length=12)
    revision_from_artifact_id: int = Field(default=0, ge=0)


class DocumentFactoryReviewRequest(BaseModel):
    artifact_id: int = Field(gt=0)


class DocumentFactoryRegenerateRequest(BaseModel):
    artifact_id: int = Field(gt=0)
    approval_id: str = Field(default="", max_length=120)
    reviewer_note: str = Field(default="", max_length=4_000)
    revision_focus: list[str] = Field(default_factory=list, max_length=8)
    acceptance_criteria: list[str] = Field(default_factory=list, max_length=12)

class DocumentDeliveryRequest(BaseModel):
    artifact_id: int
    formats: list[str] = Field(default_factory=lambda: ["docx", "pdf"], max_length=2)
    title: str = Field(default="", max_length=240)
    parent_approval_id: str = Field(default="", max_length=120)

def document_factory_source_descriptors(limit: int = 100) -> list[dict[str, Any]]:
    projects = {str(item.get("id")): item for item in load_projects()}
    descriptors = []
    seen_paths: set[str] = set()
    for artifact in _app_call('_app_call', 'list_artifacts', ):
        source_key = str(artifact.get("path") or f"artifact:{artifact.get('id')}")
        if source_key in seen_paths:
            continue
        seen_paths.add(source_key)
        path, path_error = _app_call('_app_call', 'artifact_source_path', artifact)
        metadata = artifact.get("metadata") if isinstance(artifact.get("metadata"), dict) else {}
        project_id = str(artifact.get("project_id") or "")
        descriptors.append(
            {
                "id": artifact.get("id"),
                "project_id": project_id,
                "project_name": agent_display_name(project_id) or projects.get(project_id, {}).get("title", project_id),
                "name": artifact.get("name", "未命名 Artifact"),
                "kind": artifact.get("kind", "file"),
                "created_at": artifact.get("created_at", ""),
                "title": metadata.get("title") or artifact.get("name", ""),
                "version": metadata.get("version"),
                "source_name": metadata.get("source_name") or artifact.get("name", ""),
                "readable": bool(path),
                "unavailable_reason": path_error,
            }
        )
    return descriptors[: max(1, min(limit, 200))]


def collect_document_factory_materials(request: DocumentFactoryRequest) -> dict[str, Any]:
    requested_ids = list(dict.fromkeys(int(value) for value in request.artifact_ids))
    materials: list[dict[str, Any]] = []
    errors: list[str] = []
    for artifact_id in requested_ids:
        artifact = _app_call('_app_call', 'get_artifact_record', artifact_id)
        if not artifact:
            errors.append(f"Artifact #{artifact_id} 不存在")
            continue
        content, error = _app_call('_app_call', 'read_artifact_source', artifact)
        if error:
            errors.append(f"{artifact.get('name', f'Artifact #{artifact_id}')}：{error}")
            continue
        metadata = artifact.get("metadata") if isinstance(artifact.get("metadata"), dict) else {}
        materials.append(
            {
                "artifact_id": artifact_id,
                "project_id": artifact.get("project_id", ""),
                "project_name": agent_display_name(str(artifact.get("project_id") or "")),
                "name": artifact.get("name", f"Artifact #{artifact_id}"),
                "kind": artifact.get("kind", "file"),
                "version": metadata.get("version"),
                "source_name": metadata.get("source_name") or artifact.get("name", ""),
                "content": content,
            }
        )
    if request.source_text.strip():
        materials.insert(
            0,
            {
                "artifact_id": None,
                "project_id": "inline",
                "project_name": "本次输入",
                "name": request.source_name.strip() or "粘贴材料",
                "kind": "inline_material",
                "version": None,
                "source_name": request.source_name.strip() or "粘贴材料",
                "content": request.source_text.strip(),
            },
        )
    combined_parts = []
    for index, material in enumerate(materials, start=1):
        artifact_label = f"Artifact #{material['artifact_id']}" if material.get("artifact_id") else "本次输入"
        version_label = f" · v{material['version']}" if material.get("version") else ""
        combined_parts.append(
            f"### 来源 {index} · {artifact_label} · {material['project_name']} · {material['name']}{version_label}\n"
            f"来源文件：{material['source_name']}\n\n{clip_for_llm(material['content'], 24_000)}"
        )
    return {
        "materials": materials,
        "errors": errors,
        "combined_text": clip_for_llm("\n\n".join(combined_parts), 80_000),
        "source_artifact_ids": [item["artifact_id"] for item in materials if item.get("artifact_id")],
    }


def document_factory_citation_coverage(document_text: str, source_materials: dict[str, Any]) -> dict[str, Any]:
    """Estimate paragraph-level source coverage without pretending to prove facts.

    The document Agent is asked to retain source markers, but a marker count is
    too weak: one marker can sit beside several unsupported paragraphs. This
    lightweight check compares lexical evidence against the registered source
    materials and reports a review signal. It never edits the document or
    upgrades an unsupported claim into a fact.
    """
    materials = source_materials.get("materials") if isinstance(source_materials, dict) else []
    source_tokens = set()
    for material in materials or []:
        if isinstance(material, dict):
            source_tokens.update(knowledge_tokens(material.get("content", "")))
    paragraphs = []
    for raw in re.split(r"\n\s*\n", str(document_text or "")):
        text = re.sub(r"^\s{0,3}(?:#{1,6}|[-*+]|•|\d+[.)])\s*", "", raw.strip(), flags=re.MULTILINE)
        text = re.sub(r"\[[^\]]{1,160}\]", "", text).strip()
        tokens = knowledge_tokens(text)
        if len(text) >= 18 and len(tokens) >= 2:
            marker = bool(re.search(r"\[\s*来源[：:]", raw, flags=re.IGNORECASE))
            overlap = len(tokens & source_tokens) / max(1, len(tokens)) if source_tokens else 0.0
            paragraphs.append({"text": clip(text, 180), "marked": marker, "overlap": round(overlap, 3), "supported": overlap >= 0.18})
    supported = [item for item in paragraphs if item["supported"]]
    marked_supported = [item for item in supported if item["marked"]]
    return {
        "paragraph_count": len(paragraphs),
        "source_supported_paragraphs": len(supported),
        "marked_supported_paragraphs": len(marked_supported),
        "coverage": round(len(marked_supported) / len(supported), 3) if supported else None,
        "marked_paragraphs": sum(1 for item in paragraphs if item["marked"]),
        "unmarked_supported_examples": [item["text"] for item in supported if not item["marked"]][:5],
        "policy": "基于词汇重叠和来源标记的人工复核提示，不是事实一致性证明。",
    }


def document_factory_review_checks(document_text: str, metadata: dict[str, Any], source_materials: dict[str, Any]) -> dict[str, Any]:
    source_ids = [int(value) for value in metadata.get("source_artifact_ids", []) if str(value).isdigit()]
    source_count = len(source_ids)
    source_markers = len(re.findall(r"\[来源[：:]", document_text))
    sensitive_hits = sorted(
        set(
            match.group(0).lower()
            for match in re.finditer(r"(?:api[_ -]?key|authorization|bearer\s+[a-z0-9._-]+|password|cookie|secret)", document_text, flags=re.IGNORECASE)
        )
    )
    citation_coverage = _app_call('_app_call', 'document_factory_citation_coverage', document_text, source_materials)
    coverage = citation_coverage.get("coverage")
    coverage_status = "warn" if source_count and (coverage is None or float(coverage) < 0.6) else "pass" if source_count else "warn"
    coverage_detail = (
        f"来源支持段落 {citation_coverage['source_supported_paragraphs']} 段，已带来源标记 {citation_coverage['marked_supported_paragraphs']} 段，覆盖率 {float(coverage) * 100:.0f}%"
        if coverage is not None
        else "没有足够的来源词汇重叠来评估段落覆盖，建议人工核对"
    )
    checks = [
        {
            "id": "document_nonempty",
            "label": "产物内容可读取",
            "status": "pass" if document_text.strip() else "fail",
            "detail": "已读取生成的 Markdown" if document_text.strip() else "产物为空",
        },
        {
            "id": "source_manifest",
            "label": "来源清单存在",
            "status": "pass" if source_count else "warn",
            "detail": f"登记了 {source_count} 份 Artifact 来源" if source_count else "这份产物没有登记 Artifact 来源",
        },
        {
            "id": "citation_coverage",
            "label": "段落级引用覆盖",
            "status": coverage_status,
            "detail": coverage_detail if source_count else "没有登记来源，无法进行段落级引用覆盖检查",
        },
        {
            "id": "sensitive_scan",
            "label": "敏感信息扫描",
            "status": "fail" if sensitive_hits else "pass",
            "detail": f"发现疑似敏感词：{'、'.join(sensitive_hits)}" if sensitive_hits else "未发现常见 Key、Cookie、密码或授权头模式",
        },
        {
            "id": "source_readability",
            "label": "来源文件仍可读取",
            "status": "pass" if source_count and not source_materials.get("errors") else "warn" if source_count else "warn",
            "detail": "来源文件均可重新读取" if source_count and not source_materials.get("errors") else "部分来源已不可读取，无法完成完整复核" if source_materials.get("errors") else "没有可复核的 Artifact 来源",
        },
    ]
    return {
        "checks": checks,
        "source_artifact_ids": source_ids,
        "source_markers": source_markers,
        "citation_coverage": citation_coverage,
        "sensitive_hits": sensitive_hits,
        "errors": [item["detail"] for item in checks if item["status"] == "fail"],
        "warnings": [item["detail"] for item in checks if item["status"] == "warn"],
    }

@app.post("/api/doc-factory/extract")
async def extract_document(upload: UploadFile = File(...)) -> dict[str, Any]:
    content, filename = await extract_upload_text(upload)
    if not content.strip():
        raise HTTPException(400, "文件中没有可提取的文本")
    return {
        "filename": filename,
        "content": clip(content, 100_000),
        "extractor": _app_call('_app_call', 'document_extraction_engine', filename),
        "markitdown": _app_call('_app_call', 'markitdown_status', ),
        "mineru": _app_call('_app_call', 'mineru_status', ),
    }


@app.get("/api/doc-factory/templates")
async def get_document_factory_templates() -> dict[str, Any]:
    return {"templates": _app_call('_app_call', 'document_factory_templates', )}


@app.get("/api/doc-factory/sources")
def get_document_factory_sources() -> dict[str, Any]:
    return {
        "sources": _app_call('_app_call', 'document_factory_source_descriptors', limit=120),
        "policy": "只读取已登记且位于工作台 outputs、knowledge-base、Obsidian Vault 或热点/行情/服务器安全快照内的文件；账户和配置快照不进入文档材料，原始文件保持只读。支持可选 MarkItDown 增强 PDF、DOCX、XLSX、PPTX、HTML 提取，未安装时回退到内置解析器。",
        "extractor": _app_call('_app_call', 'markitdown_status', ),
        "mineru": _app_call('_app_call', 'mineru_status', ),
    }


def document_factory_history(artifact_id: int = 0, title: str = "") -> list[dict[str, Any]]:
    """Return the complete local version chain for a document title."""
    artifacts = _app_call('_app_call', 'list_artifacts', "doc-factory")
    selected = _app_call('_app_call', 'get_artifact_record', artifact_id) if artifact_id else None
    selected_metadata = selected.get("metadata") if isinstance(selected, dict) else {}
    document_title = title.strip() or str(selected_metadata.get("title") or "")
    if not document_title and selected:
        document_title = str(selected.get("name") or "")
    items = []
    for artifact in artifacts:
        metadata = artifact.get("metadata") if isinstance(artifact.get("metadata"), dict) else {}
        if document_title and str(metadata.get("title") or "") != document_title:
            continue
        if not document_title and selected and artifact.get("id") != selected.get("id"):
            continue
        items.append({
            "id": artifact.get("id"),
            "name": artifact.get("name", ""),
            "kind": artifact.get("kind", ""),
            "path": artifact.get("path", ""),
            "version": metadata.get("version"),
            "title": metadata.get("title", ""),
            "created_at": artifact.get("created_at", ""),
            "previous_artifact_id": metadata.get("previous_artifact_id"),
            "citation_coverage": metadata.get("citation_coverage", {}),
            "warnings": metadata.get("warnings", []),
        })
    return sorted(items, key=lambda item: (int(item.get("version") or 0), str(item.get("created_at") or "")), reverse=True)[:100]


@app.get("/api/doc-factory/history")
def get_document_factory_history(artifact_id: int = 0, title: str = "") -> dict[str, Any]:
    if artifact_id and not _app_call('_app_call', 'get_artifact_record', artifact_id):
        raise HTTPException(404, "文档 Artifact 不存在")
    history = _app_call('_app_call', 'document_factory_history', artifact_id, title)
    return {"history": history, "count": len(history), "policy": "版本只读来自 Artifact 与 version_of 关系；重新生成会创建新版本，不覆盖旧文件。"}


@app.post("/api/doc-factory/validate")
def validate_document_factory(request: DocumentFactoryRequest) -> dict[str, Any]:
    materials = _app_call('_app_call', 'collect_document_factory_materials', request)
    return {"validation": _app_call('_app_call', 'validate_document_factory_payload', request, materials)}


@app.post("/api/doc-factory/run")
async def run_document_factory(request: DocumentFactoryRequest) -> dict[str, Any]:
    if not _app_call('_app_call', 'llm_settings', )["configured"]:
        raise HTTPException(503, "请先在工作台顶部配置全局 LLM")
    materials = await asyncio.to_thread(_app_call, 'collect_document_factory_materials', request)
    validation = await asyncio.to_thread(_app_call, 'validate_document_factory_payload', request, materials)
    if not validation["valid"]:
        raise HTTPException(400, "；".join(validation["errors"]))
    template = _app_call('_app_call', '_DOC_FACTORY_TEMPLATES', )[request.template]
    previous_artifact = next(
        (
            artifact
            for artifact in _app_call('_app_call', 'list_artifacts', "doc-factory")
            if artifact.get("metadata", {}).get("title") == request.title.strip()
        ),
        None,
    )
    previous_version = int((previous_artifact or {}).get("metadata", {}).get("version") or 0)
    version = previous_version + 1
    revision_focus = [clip(str(item).strip(), 120) for item in request.revision_focus if str(item).strip()][:8]
    acceptance_criteria = [clip(str(item).strip(), 240) for item in request.acceptance_criteria if str(item).strip()][:12]
    revision_brief = ""
    if revision_focus or acceptance_criteria:
        revision_brief = (
            "\n本轮为结构化修订，请优先处理以下修订重点："
            f"{'、'.join(revision_focus) or '按审批意见'}。"
            f"验收标准：{'；'.join(acceptance_criteria) or '保留原有事实边界并补齐来源标记'}。"
        )
    system = (
        "你是本地文档工厂。根据用户指令处理输入材料，输出可直接交付的中文 Markdown。"
        "不要编造材料中没有的信息；如果信息不足，请明确标注。"
        f"当前模板是「{template['label']}」：{template['description']}"
        f"{revision_brief}"
    )
    prompt = (
        f"文档模板：{template['label']}\n模板目标：{template['instruction']}\n\n"
        f"用户补充要求：\n{request.instruction}\n\n"
        f"结构化修订重点：{'、'.join(revision_focus) or '本轮初稿，无额外修订重点'}\n"
        f"验收标准：{'；'.join(acceptance_criteria) or '结论、事实、判断、风险和下一步边界清楚；关键内容保留来源标记'}\n\n"
        "以下材料按来源分段。请在关键事实、数据和判断后保留来源标记（例如 [来源：Artifact #12 · 行情快照]），"
        "不确定或来源冲突处明确标注‘待核实’，不要把不同来源的事实拼成一个未经证实的结论。\n\n"
        f"材料：\n{materials['combined_text']}"
    )
    try:
        answer = await _app_call('_app_call', 'call_llm', [
            {"role": "system", "content": system},
            {"role": "user", "content": prompt},
        ])
    except httpx.HTTPStatusError as exc:
        detail = clip(exc.response.text, 500)
        raise HTTPException(502, f"文档生成失败：上游返回 {exc.response.status_code}：{detail}") from exc
    except Exception as exc:
        raise HTTPException(502, f"文档生成失败：{exc}") from exc
    output_name = f"{datetime.now().strftime('%Y%m%d-%H%M%S')}-v{version}-{_app_call('_app_call', 'safe_filename', request.title, '文档产物')}.md"
    output_path = OUTPUTS_DIR / output_name
    output_path.write_text(answer.rstrip() + "\n", encoding="utf-8")
    citation_coverage = _app_call('_app_call', 'document_factory_citation_coverage', answer, materials)
    artifact = await asyncio.to_thread(_app_call, 'register_artifact_safely', 
        project_id="doc-factory",
        name=output_name,
        path=str(output_path),
        kind="document_output",
        metadata={
            "title": request.title.strip(),
            "template": request.template,
            "version": version,
            "source_name": request.source_name.strip() or "粘贴材料",
            "source_chars": len(materials["combined_text"]),
            "source_artifacts": validation.get("materials", []),
            "source_artifact_ids": materials.get("source_artifact_ids", []),
            "warnings": validation["warnings"],
            "instruction": clip(request.instruction, 500),
            "revision_focus": revision_focus,
            "acceptance_criteria": acceptance_criteria,
            "revision_from_artifact_id": request.revision_from_artifact_id or None,
            "previous_artifact_id": previous_artifact.get("id") if previous_artifact else None,
            "citation_coverage": citation_coverage,
        },
    )
    relation = None
    if previous_artifact and artifact:
        relation = await asyncio.to_thread(_app_call, 'create_relation_record', 
            from_type="artifact",
            from_id=str(previous_artifact["id"]),
            to_type="artifact",
            to_id=str(artifact["id"]),
            relation_type="version_of",
            metadata={"project_id": "doc-factory", "title": request.title.strip(), "version": version},
        )
    source_relations = []
    if artifact:
        for source in validation.get("materials", []):
            source_artifact_id = source.get("artifact_id")
            if not source_artifact_id:
                continue
            source_relation = await asyncio.to_thread(_app_call, 'create_relation_record', 
                from_type="artifact",
                from_id=str(source_artifact_id),
                to_type="artifact",
                to_id=str(artifact.get("id")),
                relation_type="used_as_document_source",
                metadata={
                    "source_project": source.get("project_id", ""),
                    "source_name": source.get("name", ""),
                    "document_title": request.title.strip(),
                    "document_version": version,
                },
            )
            source_relations.append(source_relation)
        if source_relations:
            try:
                await asyncio.to_thread(_app_call, 'create_notification_record', 
                    title="文档已生成并保留来源链",
                    body=f"《{request.title.strip()}》引用了 {len(source_relations)} 份工作区 Artifact，可从文档工厂回溯来源。",
                    project_id="doc-factory",
                    kind="agent_result",
                    level="info",
                    href="/projects/doc-factory",
                    event_key=f"doc-factory:sources:{artifact.get('id')}",
                    dedupe_seconds=0,
                )
            except Exception:
                log.debug("忽略异常（run_document_factory）", exc_info=True)
    return {
        "answer": answer,
        "filename": output_name,
        "path": str(output_path),
        "artifact": artifact,
        "relation": relation,
        "source_relations": source_relations,
        "materials": validation.get("materials", []),
        "validation": validation,
    }


@app.post("/api/doc-factory/regenerate")
async def regenerate_document_factory(request: DocumentFactoryRegenerateRequest) -> dict[str, Any]:
    """Create a new document version from the previous version's sources and review note."""
    artifact = await asyncio.to_thread(_app_call, 'get_artifact_record', request.artifact_id)
    if not artifact or artifact.get("project_id") != "doc-factory":
        raise HTTPException(404, "待修改的文档产物不存在")
    metadata = artifact.get("metadata") if isinstance(artifact.get("metadata"), dict) else {}
    parent_approval_id = str(request.approval_id or "").strip()
    parent_approval_payload: dict[str, Any] = {}
    if parent_approval_id:
        connection = _app_call('_app_call', 'db_connection', )
        try:
            parent_row = connection.execute("SELECT status, kind, payload_json FROM approval_requests WHERE id = ?", (parent_approval_id,)).fetchone()
        finally:
            connection.close()
        if not parent_row or parent_row["kind"] != "document_delivery":
            raise HTTPException(404, "关联的文档审批不存在")
        if parent_row["status"] not in {"changes_requested", "rejected"}:
            raise HTTPException(409, "只有需要修改或已退回的审批才能发起修订")
        parent_approval_payload = _app_call('_app_call', 'platform_decode_json', parent_row["payload_json"], {})
    source_ids = [int(value) for value in metadata.get("source_artifact_ids", []) if str(value).isdigit()]
    previous_text, read_error = _app_call('_app_call', 'read_artifact_source', artifact)
    if read_error:
        raise HTTPException(409, f"无法读取上一版文档：{read_error}")
    instruction_parts = [str(metadata.get("instruction") or "").strip(), f"审批/复核意见：{request.reviewer_note.strip()}" if request.reviewer_note.strip() else "", "请根据以上意见生成新版本，并保留可回溯来源标记。"]
    regeneration = DocumentFactoryRequest(
        title=str(metadata.get("title") or artifact.get("name") or "修改后的文档"),
        source_text="" if source_ids else clip(previous_text, 100_000),
        instruction="\n\n".join(part for part in instruction_parts if part),
        template=str(metadata.get("template") or "general_report"),
        source_name=str(metadata.get("source_name") or "上一版文档"),
        artifact_ids=source_ids,
        revision_focus=request.revision_focus or [str(item) for item in metadata.get("revision_focus", []) if str(item)],
        acceptance_criteria=request.acceptance_criteria or [str(item) for item in metadata.get("acceptance_criteria", []) if str(item)],
        revision_from_artifact_id=request.artifact_id,
    )
    result = await _app_call('_app_call', 'run_document_factory', regeneration)
    result["regenerated_from_artifact_id"] = request.artifact_id
    new_artifact_id = int((result.get("artifact") or {}).get("id") or 0)
    if new_artifact_id:
        formats = [str(value).lower().strip() for value in (parent_approval_payload.get("formats") or ["docx"]) if str(value).lower().strip() in {"docx", "pdf"}]
        try:
            delivery = await asyncio.to_thread(_app_call, 'deliver_document_factory', DocumentDeliveryRequest(artifact_id=new_artifact_id, formats=formats or ["docx"], title=str(metadata.get("title") or artifact.get("name") or "Workbench 文档"), parent_approval_id=parent_approval_id))
            result["revision_delivery"] = delivery
            result["message"] = "新版本已生成，并已创建新的交付审批轮次。"
        except HTTPException as exc:
            # The Markdown revision is still useful if an optional PDF
            # converter is unavailable; surface the exact next action.
            result["revision_delivery_error"] = str(exc.detail)
            result["message"] = f"新版本已生成，但交付包创建失败：{exc.detail}"
    return result


@app.post("/api/doc-factory/review")
async def review_document_factory_output(request: DocumentFactoryReviewRequest) -> dict[str, Any]:
    if not _app_call('_app_call', 'llm_settings', )["configured"]:
        raise HTTPException(503, "请先在工作台顶部配置全局 LLM")
    artifact = await asyncio.to_thread(_app_call, 'get_artifact_record', request.artifact_id)
    if not artifact or artifact.get("project_id") != "doc-factory":
        raise HTTPException(404, "文档产物 Artifact 不存在")
    document_text, read_error = _app_call('_app_call', 'read_artifact_source', artifact)
    if read_error:
        raise HTTPException(409, f"无法读取待校验文档：{read_error}")
    metadata = artifact.get("metadata") if isinstance(artifact.get("metadata"), dict) else {}
    source_ids = [int(value) for value in metadata.get("source_artifact_ids", []) if str(value).isdigit()]
    source_materials = await asyncio.to_thread(_app_call, 'collect_document_factory_materials', DocumentFactoryRequest(artifact_ids=source_ids))
    checks = _app_call('_app_call', 'document_factory_review_checks', document_text, metadata, source_materials)
    run = await asyncio.to_thread(_app_call, 'create_agent_run_record', 
        project_id="doc-factory",
        kind="document_review",
        title=f"校验文档：{metadata.get('title') or artifact.get('name', '未命名产物')}",
        request={"artifact_id": request.artifact_id, "source_artifact_ids": source_ids},
        max_attempts=2,
    )
    add_agent_run_event(run["id"], "review_started", "开始执行来源、引用和敏感信息二次校验。", metadata={"artifact_id": request.artifact_id})
    source_context = source_materials.get("combined_text") or "（没有可重新读取的来源 Artifact，仅能做格式和敏感信息检查。）"
    prompt = (
        "请作为文档校验 Agent，审查下面的生成文档是否忠实于来源材料。只指出证据支持、证据不足、来源冲突、"
        "引用缺失和需要人工确认的地方，不要替用户改写原文，不要编造新事实。输出简洁中文 Markdown，包含：\n"
        "1. 事实一致性；2. 引用覆盖；3. 冲突或疑点；4. 修改建议。\n\n"
        f"来源材料：\n{source_context}\n\n待校验文档：\n{clip_for_llm(document_text, 28_000)}"
    )
    try:
        agent_answer = await _app_call('_app_call', 'call_llm', 
            [
                {"role": "system", "content": "你是保守的事实与引用校验 Agent。无法从来源确认的内容必须标记为待核实。"},
                {"role": "user", "content": prompt},
            ],
            max_tokens=3000,
            temperature=0.1,
        )
    except httpx.HTTPStatusError as exc:
        detail = clip(exc.response.text, 500)
        error = f"文档校验失败：上游返回 {exc.response.status_code}：{detail}"
        await asyncio.to_thread(_app_call, 'update_agent_run_record', run["id"], status="failed", error=error)
        await asyncio.to_thread(_app_call, 'add_agent_run_event', run["id"], "review_failed", error, level="error")
        raise HTTPException(502, error) from exc
    except Exception as exc:
        error = f"文档校验失败：{exc}"
        await asyncio.to_thread(_app_call, 'update_agent_run_record', run["id"], status="failed", error=error)
        await asyncio.to_thread(_app_call, 'add_agent_run_event', run["id"], "review_failed", error, level="error")
        raise HTTPException(502, error) from exc
    review_title = metadata.get("title") or artifact.get("name", "未命名产物")
    review_name = f"{datetime.now().strftime('%Y%m%d-%H%M%S')}-校验-{_app_call('_app_call', 'safe_filename', review_title, '文档产物')}.md"
    review_path = OUTPUTS_DIR / review_name
    check_lines = "\n".join(
        f"- {'✅' if item['status'] == 'pass' else '⚠️' if item['status'] == 'warn' else '❌'} **{item['label']}**：{item['detail']}"
        for item in checks["checks"]
    )
    review_body = (
        f"# 文档校验报告：{review_title}\n\n"
        f"> 被校验产物：{artifact.get('name', '')}（Artifact #{artifact.get('id')}）\n"
        f"> 校验时间：{now_iso()}\n\n"
        "## 自动检查\n\n"
        f"{check_lines}\n\n"
        "## Agent 二次审查\n\n"
        f"{agent_answer.strip()}\n"
    )
    review_path.write_text(review_body.rstrip() + "\n", encoding="utf-8")
    review_artifact = await asyncio.to_thread(_app_call, 'register_artifact_safely', 
        project_id="doc-factory",
        name=review_name,
        path=str(review_path),
        kind="document_review",
        metadata={
            "title": f"{review_title} · 校验报告",
            "reviewed_artifact_id": artifact.get("id"),
            "source_artifact_ids": source_ids,
            "check_status": "fail" if checks["errors"] else "warn" if checks["warnings"] else "pass",
            "checks": checks["checks"],
            "run_id": run["id"],
        },
    )
    relations = []
    if review_artifact:
        relations.append(
            _app_call('_app_call', 'create_relation_record', 
                from_type="artifact",
                from_id=str(artifact.get("id")),
                to_type="artifact",
                to_id=str(review_artifact.get("id")),
                relation_type="reviewed_by",
                metadata={"run_id": run["id"], "status": "fail" if checks["errors"] else "warn" if checks["warnings"] else "pass"},
            )
        )
        for source_id in source_ids:
            relations.append(
                _app_call('_app_call', 'create_relation_record', 
                    from_type="artifact",
                    from_id=str(source_id),
                    to_type="artifact",
                    to_id=str(review_artifact.get("id")),
                    relation_type="review_source",
                    metadata={"reviewed_artifact_id": artifact.get("id"), "run_id": run["id"]},
                )
            )
    result = {
        "artifact": review_artifact,
        "reviewed_artifact": artifact,
        "checks": checks,
        "answer": agent_answer,
        "path": str(review_path),
        "run_id": run["id"],
        "relations": relations,
    }
    await asyncio.to_thread(_app_call, 'update_agent_run_record', run["id"], status="succeeded", result={"review_artifact_id": review_artifact.get("id") if review_artifact else None, "checks": checks})
    add_agent_run_event(run["id"], "review_succeeded", "文档校验报告已保存。", metadata={"review_artifact_id": review_artifact.get("id") if review_artifact else None})
    try:
        await asyncio.to_thread(_app_call, 'create_notification_record', 
            title="文档校验完成",
            body=f"《{review_title}》的事实/引用校验报告已生成，可查看自动检查和 Agent 二次审查。",
            project_id="doc-factory",
            kind="agent_result",
            level="warning" if checks["errors"] or checks["warnings"] else "info",
            href="/projects/doc-factory",
            event_key=f"doc-factory:review:{review_artifact.get('id') if review_artifact else run['id']}",
            dedupe_seconds=0,
        )
    except Exception:
        log.debug("忽略异常（review_document_factory_output）", exc_info=True)
    return result

@app.post("/api/doc-factory/deliver")
def deliver_document_factory(request: DocumentDeliveryRequest) -> dict[str, Any]:
    artifact = _app_call('_app_call', 'get_artifact_record', request.artifact_id)
    if not artifact or artifact.get("project_id") != "doc-factory":
        raise HTTPException(404, "文档产物 Artifact 不存在")
    text, read_error = _app_call('_app_call', 'read_artifact_source', artifact)
    if read_error:
        raise HTTPException(409, f"无法读取文档产物：{read_error}")
    metadata = artifact.get("metadata") if isinstance(artifact.get("metadata"), dict) else {}
    parent_approval_id = str(request.parent_approval_id or "").strip()
    parent_payload: dict[str, Any] = {}
    parent_round = 0
    if parent_approval_id:
        connection = _app_call('_app_call', 'db_connection', )
        try:
            parent_row = connection.execute("SELECT payload_json FROM approval_requests WHERE id = ? AND kind = 'document_delivery'", (parent_approval_id,)).fetchone()
        finally:
            connection.close()
        if parent_row:
            parent_payload = _app_call('_app_call', 'platform_decode_json', parent_row["payload_json"], {})
            try:
                parent_round = max(0, int(parent_payload.get("round") or 0))
            except (TypeError, ValueError):
                parent_round = 0
    approval_round = parent_round + 1 if parent_approval_id else max(1, int(metadata.get("approval_round") or 1))
    previous_artifact_id = metadata.get("revision_from_artifact_id")
    source_artifact_id = previous_artifact_id if parent_approval_id and previous_artifact_id else artifact.get("id")
    revision_artifact_id = artifact.get("id") if parent_approval_id and previous_artifact_id else None
    title = request.title.strip() or str(metadata.get("title") or artifact.get("name") or "Workbench 文档")
    base = _app_call('_app_call', 'safe_filename', title, "workbench-document")
    version = int(metadata.get("version") or 1)
    created: list[dict[str, Any]] = []
    formats = [value.lower().strip() for value in request.formats if value.lower().strip() in {"docx", "pdf"}]
    if not formats:
        raise HTTPException(400, "formats 只能是 docx 或 pdf")
    docx_path = OUTPUTS_DIR / f"{base}-v{version}.docx"
    if "docx" in formats or "pdf" in formats:
        _app_call('_app_call', 'build_docx_delivery', title, text, docx_path)
        docx_artifact = _app_call('_app_call', 'register_artifact_safely', project_id="doc-factory", name=docx_path.name, path=str(docx_path), kind="document_delivery_docx", metadata={"title": title, "version": version, "source_artifact_id": artifact.get("id"), "revision_artifact_id": revision_artifact_id, "format": "docx", "approval_status": "pending", "approval_round": approval_round, "parent_approval_id": parent_approval_id})
        created.append(docx_artifact or {"name": docx_path.name, "path": str(docx_path)})
        _app_call('_app_call', 'create_relation_record', from_type="artifact", from_id=str(artifact.get("id")), to_type="artifact", to_id=str((docx_artifact or {}).get("id", "")), relation_type="delivered_as_docx", metadata={"title": title, "version": version}) if docx_artifact else None
    if "pdf" in formats:
        pdf_path = OUTPUTS_DIR / f"{base}-v{version}.pdf"
        ok, conversion_error = _app_call('_app_call', 'convert_docx_to_pdf', docx_path, pdf_path)
        if not ok:
            raise HTTPException(503, f"PDF 交付需要 LibreOffice：{conversion_error}")
        pdf_artifact = _app_call('_app_call', 'register_artifact_safely', project_id="doc-factory", name=pdf_path.name, path=str(pdf_path), kind="document_delivery_pdf", metadata={"title": title, "version": version, "source_artifact_id": artifact.get("id"), "revision_artifact_id": revision_artifact_id, "format": "pdf", "approval_status": "pending", "approval_round": approval_round, "parent_approval_id": parent_approval_id})
        created.append(pdf_artifact or {"name": pdf_path.name, "path": str(pdf_path)})
        _app_call('_app_call', 'create_relation_record', from_type="artifact", from_id=str(artifact.get("id")), to_type="artifact", to_id=str((pdf_artifact or {}).get("id", "")), relation_type="delivered_as_pdf", metadata={"title": title, "version": version}) if pdf_artifact else None
    approval_payload = {
        "round": approval_round,
        "source_artifact_id": source_artifact_id,
        "current_artifact_id": artifact.get("id"),
        "revision_artifact_id": revision_artifact_id,
        "parent_approval_id": parent_approval_id,
        "delivery_artifacts": [item.get("id") for item in created if item.get("id")],
        "formats": formats,
        "title": title,
        "version": version,
    }
    approval = _app_call('_app_call', 'create_approval_request', "doc-factory", "document_delivery", f"第 {approval_round} 轮审批：{title}", approval_payload)
    if parent_approval_id:
        _app_call('_app_call', 'create_relation_record', from_type="approval", from_id=parent_approval_id, to_type="approval", to_id=approval["id"], relation_type="approval_revised_as", metadata={"round": approval_round, "source_artifact_id": source_artifact_id, "revision_artifact_id": revision_artifact_id})
    create_notification_record(title="正式文档交付包待审批", body=f"{title} · {', '.join(formats).upper()} · 请在审批中心确认或提出修改意见。", project_id="doc-factory", kind="approval", level="warning", href="/projects/doc-factory", event_key=f"document-delivery:{approval['id']}", dedupe_seconds=0)
    return {"ok": True, "source": artifact, "deliveries": created, "approval": approval, "approval_context": {"round": approval_round, "parent_approval_id": parent_approval_id, "source_artifact_id": source_artifact_id, "revision_artifact_id": revision_artifact_id}, "message": f"第 {approval_round} 轮交付包已生成，当前状态为待审批。"}


DOC_FACTORY_TEMPLATES: dict[str, dict[str, Any]] = {
    "general_report": {
        "label": "通用分析报告",
        "description": "结论、事实、风险和下一步，适合大多数材料；例如把调研资料整理成一份可交付的分析。",
        "instruction": "整理成结构清晰的中文 Markdown，先给出结论，再列出关键事实、风险和下一步行动。",
    },
    "meeting_notes": {
        "label": "会议纪要",
        "description": "议题、决策、待办、负责人和截止时间；适合把会议记录或语音转写整理成纪要。",
        "instruction": "整理成会议纪要：会议结论、关键讨论、已确认决策、待办事项（负责人/截止时间）和待确认问题。不要补造材料中没有的负责人或日期。",
    },
    "prd": {
        "label": "产品需求文档",
        "description": "背景、目标、用户、范围、流程、指标和风险；适合把一个想法或讨论整理成可评审的 PRD。",
        "instruction": "整理成产品需求文档：背景与问题、目标、不做什么、目标用户、核心流程、功能需求、数据指标、验收标准和风险。材料缺失处明确标注待补充。",
    },
    "weekly_brief": {
        "label": "周报/简报",
        "description": "本周进展、关键变化、风险和下周计划；适合把零散工作记录汇总成周报。",
        "instruction": "整理成中文周报：本周完成、关键数据或事实、问题与风险、需要协同的事项、下周计划。每项结论都尽量保留材料来源。",
    },
    "action_list": {
        "label": "行动清单",
        "description": "把材料转成可执行任务和确认项；适合把长文或会议内容拆成待办。",
        "instruction": "整理成行动清单，包含事项、负责人（如果材料中有）、截止时间（如果材料中有）、优先级、依赖和当前状态；缺失字段写待确认。",
    },
    "study_notes": {
        "label": "学习笔记/知识卡片",
        "description": "把课程、文章或资料整理成可复习的概念卡片、例子和练习；适合持续沉淀到知识库。",
        "instruction": "整理成学习笔记：先写一句话摘要，再列核心概念、概念之间的关系、关键例子、容易混淆的点、待验证问题和复习题。材料缺失处明确标注待补充，不要编造结论或来源。",
    },
    "decision_record": {
        "label": "决策记录",
        "description": "记录背景、选项、取舍和后续验证；适合保存产品、技术和个人工作决策。",
        "instruction": "整理成决策记录：背景与问题、目标、约束、可选方案、比较依据、最终决定、明确不选什么、风险与假设、后续验证动作和复盘时间。材料缺失处明确标注待补充，不要把推测写成已确认事实。",
    },
}


def document_factory_templates() -> list[dict[str, str]]:
    return [{"id": key, **value} for key, value in DOC_FACTORY_TEMPLATES.items()]


def validate_document_factory_payload(request: DocumentFactoryRequest, materials: dict[str, Any] | None = None) -> dict[str, Any]:
    errors: list[str] = []
    warnings: list[str] = []
    materials = materials or _app_call('collect_document_factory_materials', request)
    template = DOC_FACTORY_TEMPLATES.get(request.template)
    if not template:
        errors.append(f"未知文档模板：{request.template}")
    if not request.title.strip():
        errors.append("产物名称不能为空")
    errors.extend(materials.get("errors", []))
    combined_text = str(materials.get("combined_text") or "")
    if not combined_text:
        errors.append("至少提供一段材料或选择一份可读取的工作区 Artifact")
    elif len(combined_text.strip()) < 80:
        warnings.append("材料少于 80 个字符，生成结果可能只能形成结构草稿")
    if not request.instruction.strip():
        errors.append("处理要求不能为空")
    if len(combined_text) > 80_000:
        warnings.append("材料超过单轮建议长度，LLM 会压缩中间内容，生成前后请核对事实")
    if request.source_text.strip() and request.source_name.strip() == "粘贴材料":
        warnings.append("尚未记录原始文件名；如来自文件，建议保留文件名便于追溯")
    if len(materials.get("materials", [])) > 1:
        warnings.append(f"本次会合并 {len(materials['materials'])} 份材料；生成结果应保留来源，不要把不同来源事实混为一谈")
    return {
        "valid": not errors,
        "errors": errors,
        "warnings": warnings,
        "template": {"id": request.template, **template} if template else {"id": request.template},
        "checks": [
            {"id": "title", "label": "产物名称", "status": "pass" if request.title.strip() else "fail"},
            {"id": "source", "label": "材料来源", "status": "pass" if combined_text else "fail"},
            {"id": "provenance", "label": f"来源可追溯（{len(materials.get('materials', []))} 份）", "status": "pass" if not materials.get("errors") and combined_text else "fail"},
            {"id": "instruction", "label": "处理要求", "status": "pass" if request.instruction.strip() else "fail"},
            {"id": "template", "label": "文档模板", "status": "pass" if template else "fail"},
        ],
        "materials": [
            {
                "artifact_id": item.get("artifact_id"),
                "project_id": item.get("project_id"),
                "project_name": item.get("project_name"),
                "name": item.get("name"),
                "kind": item.get("kind"),
                "version": item.get("version"),
                "source_name": item.get("source_name"),
            }
            for item in materials.get("materials", [])
        ],
    }


def set_docx_run_font(run: Any, name: str = "Hiragino Sans GB", size: float = 11, color: str = "1F2937", bold: bool | None = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_docx_cell_shading(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def build_docx_delivery(title: str, text: str, target: Path) -> None:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = document.styles["Normal"]
    normal.font.name = "Hiragino Sans GB"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (("Heading 1", 16, "2E74B5", 16, 8), ("Heading 2", 13, "2E74B5", 12, 6), ("Heading 3", 12, "1F4D78", 8, 4)):
        style = document.styles[name]
        style.font.name = "Hiragino Sans GB"
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Hiragino Sans GB")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Hiragino Sans GB")
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
        style._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Hiragino Sans GB")
        style.font.size = Pt(size)
        style.font.color.rgb = __import__("docx").shared.RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1
    header = section.header.paragraphs[0]
    header.text = "Workbench · 文档交付包"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_docx_run_font(run, size=9, color="6B7280")
    footer = section.footer.paragraphs[0]
    footer.text = f"{WORKBENCH_VERSION} · 生成于 {datetime.now().astimezone().strftime('%Y-%m-%d %H:%M')}"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_docx_run_font(run, size=8.5, color="6B7280")
    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_after = Pt(4)
    title_run = title_paragraph.add_run(title or "未命名文档")
    set_docx_run_font(title_run, size=23, color="0B2545", bold=True)
    meta = document.add_paragraph()
    meta.paragraph_format.space_after = Pt(14)
    meta_run = meta.add_run(f"Workbench 正式交付草稿 · 版本 {WORKBENCH_VERSION}")
    set_docx_run_font(meta_run, size=10, color="6B7280")
    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(12)
    p_pr = rule._p.get_or_add_pPr()
    from docx.oxml import OxmlElement
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D8DEE8")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    def add_markdown_runs(paragraph: Any, content: str) -> None:
        """Render the small Markdown subset accepted by the document factory."""
        chunks = re.split(r"(\*\*.+?\*\*)", str(content or ""))
        for chunk in chunks:
            if not chunk:
                continue
            is_bold = chunk.startswith("**") and chunk.endswith("**") and len(chunk) >= 4
            value = chunk[2:-2] if is_bold else chunk
            run = paragraph.add_run(value)
            set_docx_run_font(run, bold=is_bold)

    for raw in str(text or "").splitlines():
        line = raw.strip()
        if not line:
            document.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        if line in {"*", "_"}:
            continue
        if line.startswith("### "):
            document.add_paragraph(line[4:], style="Heading 3")
        elif line.startswith("## "):
            document.add_paragraph(line[3:], style="Heading 2")
        elif line.startswith("# "):
            document.add_paragraph(line[2:], style="Heading 1")
        elif line.startswith(("- ", "* ")):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.space_after = Pt(8)
            add_markdown_runs(paragraph, line[2:])
        elif re.match(r"^\d+[.)] ", line):
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.space_after = Pt(8)
            add_markdown_runs(paragraph, re.sub(r"^\d+[.)] ", "", line))
        else:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.1
            add_markdown_runs(paragraph, line)
    document.save(target)


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> tuple[bool, str]:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return False, "未找到 LibreOffice/soffice"
    temp_dir = pdf_path.parent / f".pdf-convert-{uuid.uuid4().hex}"
    temp_dir.mkdir(parents=True, exist_ok=True)
    try:
        # The bundled headless LibreOffice build does not always inherit the
        # macOS font database, so Chinese TTC fonts can otherwise become
        # missing-glyph boxes.  Give fontconfig an explicit, per-conversion
        # search path.  WORKBENCH_DOCX_FONT_DIR is also supported for Linux
        # deployments where a CJK font package is mounted separately.
        font_dirs: list[Path] = []
        configured_font_dir = os.getenv("WORKBENCH_DOCX_FONT_DIR", "").strip()
        if configured_font_dir:
            font_dirs.append(Path(configured_font_dir).expanduser())
        font_dirs.extend(
            Path(path)
            for path in (
                "/System/Library/Fonts",
                "/System/Library/Fonts/Supplemental",
                "/Library/Fonts",
                str(Path.home() / "Library" / "Fonts"),
                "/usr/share/fonts",
                "/usr/local/share/fonts",
            )
        )
        font_dirs = [path for path in font_dirs if path.exists() and path.is_dir()]
        conversion_env = os.environ.copy()
        if font_dirs:
            fontconfig_path = temp_dir / "fontconfig.conf"
            fontconfig_dirs = "".join(f"    <dir>{str(path)}</dir>\n" for path in font_dirs)
            fontconfig_path.write_text(
                "<?xml version=\"1.0\"?>\n"
                "<!DOCTYPE fontconfig SYSTEM \"urn:fontconfig:fonts.dtd\">\n"
                "<fontconfig>\n"
                f"{fontconfig_dirs}"
                "    <dir prefix=\"xdg\">fonts</dir>\n"
                "    <dir>~/.fonts</dir>\n"
                "</fontconfig>\n",
                encoding="utf-8",
            )
            conversion_env["FONTCONFIG_FILE"] = str(fontconfig_path)
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(temp_dir), str(docx_path)],
            capture_output=True,
            text=True,
            timeout=60,
            check=False,
            env=conversion_env,
        )
        converted = temp_dir / f"{docx_path.stem}.pdf"
        if result.returncode != 0 or not converted.exists():
            return False, clip(result.stderr or result.stdout or "PDF 转换失败", 500)
        shutil.copy2(converted, pdf_path)
        return True, ""
    except (OSError, subprocess.SubprocessError) as exc:
        return False, str(exc)
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


__all__ = [
    "DOC_FACTORY_TEMPLATES",
    "DocumentDeliveryRequest",
    "DocumentFactoryRegenerateRequest",
    "DocumentFactoryRequest",
    "DocumentFactoryReviewRequest",
    "_DOC_FACTORY_TEMPLATES",
    "build_docx_delivery",
    "collect_document_factory_materials",
    "convert_docx_to_pdf",
    "deliver_document_factory",
    "document_factory_citation_coverage",
    "document_factory_history",
    "document_factory_review_checks",
    "document_factory_source_descriptors",
    "document_factory_templates",
    "extract_document",
    "get_document_factory_history",
    "get_document_factory_sources",
    "get_document_factory_templates",
    "regenerate_document_factory",
    "review_document_factory_output",
    "run_document_factory",
    "set_docx_cell_shading",
    "set_docx_run_font",
    "validate_document_factory",
    "validate_document_factory_payload",
]
